"""Generate business-style audit deliverable as DOCX.

Structure:
 - Title + subtitle
 - Info table (target, date, scan grade, author)
 - 1. Vezetői összefoglaló
 - 2. Technikai megállapítások (2.1 headers, 2.2/2.3 adaptive)
 - 3. Kockázati besorolás
 - 4. Javasolt megoldások (Ajánlat 1 + 2)
 - 5. Javaslat és következő lépés
"""
from __future__ import annotations

from collections import OrderedDict
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from vbeye import phrases
from vbeye.config import Config
from vbeye.scoring import CheckerResult, Severity, compute_score, grade_from_score


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str)


def _set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            spec = kwargs[edge]
            elem = OxmlElement(f"w:{edge}")
            elem.set(qn("w:val"), spec.get("val", "single"))
            elem.set(qn("w:sz"), str(spec.get("sz", 4)))
            elem.set(qn("w:space"), "0")
            elem.set(qn("w:color"), spec.get("color", "auto"))
            tcBorders.append(elem)
    tcPr.append(tcBorders)


def _set_paragraph_bottom_border(paragraph, color: str = "00B4D8", size: int = 12) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_run(paragraph, text: str, *, bold: bool = False, color: str = "2B2B2B", size: int | None = None):
    run = paragraph.add_run(text)
    run.font.color.rgb = _rgb(color)
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return run


def _add_paragraph(doc, text: str, *, color: str = "2B2B2B", size: int | None = None,
                   bold: bool = False, justify: bool = True, space_after: int = 60) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(space_after / 10)
    pf.line_spacing = 1.2
    _add_run(p, text, bold=bold, color=color, size=size)


def _add_section_heading(doc, text: str, branding) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(6)
    _add_run(p, text, bold=True, color=branding.color_primary, size=14)
    _set_paragraph_bottom_border(p, color=branding.color_accent, size=12)


def _add_subsection_heading(doc, text: str, branding) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    _add_run(p, text, bold=True, color=branding.color_primary, size=12)


def _add_bullet(doc, label: str, body: str, branding) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    if label:
        _add_run(p, f"{label} ", bold=True, color=branding.color_body)
        _add_run(p, f"– {body}", color=branding.color_body)
    else:
        _add_run(p, body, color=branding.color_body)


def _add_title_block(doc, branding) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _add_run(p, "BIZTONSÁGI AUDIT JELENTÉS", bold=True, color=branding.color_primary, size=22)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    r = sub.add_run("Webhely sérülékenységi gyors-elemzés")
    r.font.color.rgb = _rgb(branding.color_accent)
    r.italic = True
    r.font.size = Pt(11)
    _set_paragraph_bottom_border(sub, color=branding.color_border, size=18)


def _add_info_table(doc, branding, *, target_url: str, scan_date: str,
                    grade: str, scan_source: str) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    col1 = Cm(4.5)
    col2 = Cm(13.5)
    for row in table.rows:
        row.cells[0].width = col1
        row.cells[1].width = col2

    grade_color = branding.color_critical if grade in ("D", "E", "F") else branding.color_primary

    rows = [
        ("Vizsgált webhely:", target_url, branding.color_body, False),
        ("Dátum:", scan_date, branding.color_body, False),
        ("Scan minősítés:", f"{grade}  ", grade_color, True),
        ("Készítette:", f"{branding.company_name} – {branding.author}", branding.color_body, False),
    ]
    for i, (label, value, color, bold_val) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c2.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # label cell
        p1 = c1.paragraphs[0]
        _add_run(p1, label, bold=True, color=branding.color_meta, size=9)
        # value cell
        p2 = c2.paragraphs[0]
        _add_run(p2, value, bold=bold_val, color=color, size=11)
        if label == "Scan minősítés:":
            ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
            _add_run(p2, f"  ({scan_source} / {ts} UTC)", color=branding.color_meta, size=8)

    # remove default borders, add only thin separator under each row
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "D0D5DD"},
                bottom={"val": "single", "sz": 4, "color": "D0D5DD"},
                left={"val": "nil", "color": "FFFFFF"},
                right={"val": "nil", "color": "FFFFFF"},
            )


def _classify_findings(results: list[CheckerResult]) -> dict:
    by_check: dict[str, list] = {}
    flags = {
        "all_headers_missing": False,
        "headers_partial": False,
        "tls_critical": False,
        "cookies_issue": False,
        "cors_issue": False,
        "disclosure": False,
        "source_issues": False,
        "headers_ok": False,
    }

    header_missing_count = 0
    header_total = 6

    for r in results:
        for f in r.findings:
            by_check.setdefault(f.check_id, []).append(f)
            if f.severity == Severity.OK:
                continue
            if f.check_id.startswith("headers.") and ".missing" in f.check_id:
                header_missing_count += 1
            if f.check_id.startswith("ssl.") and f.severity in (Severity.HIGH, Severity.CRITICAL):
                flags["tls_critical"] = True
            if f.check_id == "headers.cookie.flags":
                flags["cookies_issue"] = True
            if f.check_id.startswith("headers.disclosure"):
                flags["disclosure"] = True
            if f.check_id.startswith("source.") and f.severity in (Severity.MEDIUM, Severity.HIGH, Severity.CRITICAL):
                flags["source_issues"] = True

    if header_missing_count >= header_total - 1:
        flags["all_headers_missing"] = True
    elif header_missing_count > 0:
        flags["headers_partial"] = True
    else:
        flags["headers_ok"] = True

    return {"by_check": by_check, "flags": flags, "headers_missing_count": header_missing_count}


def _executive_summary(doc, branding, *, host: str, grade: str, info: dict, industry: str, compliance: str) -> None:
    _add_section_heading(doc, "1. Vezetői összefoglaló", branding)

    intro = phrases.EXECUTIVE_SUMMARY_TEMPLATE["intro"].format(host=host, grade=grade)
    summary_keys = []
    if info["flags"]["all_headers_missing"]:
        summary_keys.append("all_headers_missing")
    elif info["flags"]["headers_partial"]:
        summary_keys.append("headers_partial")
    else:
        summary_keys.append("headers_ok")
    if info["flags"]["tls_critical"]:
        summary_keys.append("tls_critical")
    if info["flags"]["source_issues"]:
        summary_keys.append("source_issues")

    framing_parts = [intro]
    for k in summary_keys:
        framing_parts.append(phrases.EXECUTIVE_SUMMARY_TEMPLATE["intro_findings_summary"][k])
    _add_paragraph(doc, " ".join(framing_parts), color=branding.color_body)

    if grade in ("C", "D", "E", "F"):
        _add_paragraph(doc, phrases.EXECUTIVE_SUMMARY_TEMPLATE["risk_framing"], color=branding.color_body)

    industry_text = phrases.EXECUTIVE_SUMMARY_TEMPLATE["industry_template"].format(
        industry=industry, compliance=compliance
    )
    _add_paragraph(doc, industry_text, color=branding.color_body)


def _section_2_headers(doc, branding, info: dict) -> None:
    _add_subsection_heading(doc, "2.1 Hiányzó / nem megfelelő biztonsági HTTP-fejlécek", branding)

    missing_headers = OrderedDict()
    for check_id, findings in info["by_check"].items():
        if check_id.startswith("headers.") and (
            ".missing" in check_id or ".weak" in check_id or ".invalid" in check_id
        ):
            header = phrases.CHECK_ID_TO_HEADER.get(check_id)
            if header and header not in missing_headers:
                missing_headers[header] = phrases.HEADER_DESCRIPTIONS.get(header, "")

    if not missing_headers:
        _add_paragraph(doc, "A vizsgált biztonsági HTTP-fejlécek megfelelően vannak konfigurálva.", color=branding.color_body)
        return

    for header, desc in missing_headers.items():
        _add_bullet(doc, header, desc, branding)

    _add_paragraph(doc, phrases.HEADERS_SECTION_CLOSER, color=branding.color_body, justify=True)


def _section_2_adaptive(doc, branding, info: dict, results: list[CheckerResult]) -> int:
    """Add 2.2, 2.3, ... adaptive subsections. Returns next subsection number."""
    next_num = 2

    def num(label: str) -> str:
        nonlocal next_num
        next_num += 1
        return f"2.{next_num} {label}"

    if info["flags"]["cookies_issue"]:
        sec = phrases.COOKIE_SECTION
        _add_subsection_heading(doc, num(sec["title"]), branding)
        _add_paragraph(doc, sec["lead"], color=branding.color_body)
        # Detect which flags are missing across findings
        evidences = []
        for f in info["by_check"].get("headers.cookie.flags", []):
            evidences.append((f.title, f.evidence.lower() if f.evidence else ""))
        missing_kinds = set()
        for _t, ev in evidences:
            if "secure" not in ev:
                missing_kinds.add("secure")
            if "httponly" not in ev:
                missing_kinds.add("httponly")
            if "samesite" not in ev:
                missing_kinds.add("samesite")
        if not missing_kinds:
            missing_kinds = {"secure", "httponly", "samesite"}
        for k in ("secure", "httponly", "samesite"):
            if k in missing_kinds:
                _add_bullet(doc, "", sec["bullets"][k], branding)
        _add_paragraph(doc, sec["closer"], color=branding.color_body)

    tls_findings = [f for r in results if r.name == "ssl" for f in r.findings
                    if f.severity in (Severity.MEDIUM, Severity.HIGH, Severity.CRITICAL)]
    if tls_findings:
        sec = phrases.TLS_SECTION
        _add_subsection_heading(doc, num(sec["title"]), branding)
        _add_paragraph(doc, sec["lead"], color=branding.color_body)
        seen = set()
        for f in tls_findings:
            if f.check_id in seen:
                continue
            seen.add(f.check_id)
            text = sec["bullets"].get(f.check_id)
            if text:
                _add_bullet(doc, "", text, branding)
            else:
                _add_bullet(doc, f.title, f.description, branding)
        _add_paragraph(doc, sec["closer"], color=branding.color_body)

    if info["flags"]["disclosure"]:
        sec = phrases.DISCLOSURE_SECTION
        _add_subsection_heading(doc, num(sec["title"]), branding)
        _add_paragraph(doc, sec["lead"], color=branding.color_body)
        for f in info["by_check"]:
            if f.startswith("headers.disclosure."):
                for finding in info["by_check"][f]:
                    _add_bullet(doc, "", f"{finding.title} – {finding.evidence}", branding)
        _add_paragraph(doc, sec["closer"], color=branding.color_body)

    source_findings = [
        f for r in results if r.name == "source"
        for f in r.findings
        if f.severity in (Severity.LOW, Severity.MEDIUM, Severity.HIGH, Severity.CRITICAL)
    ]
    if source_findings:
        sec = phrases.SOURCE_SECTION
        _add_subsection_heading(doc, num(sec["title"]), branding)
        _add_paragraph(doc, sec["lead"], color=branding.color_body)
        seen = set()
        for f in source_findings:
            if f.check_id in seen:
                continue
            seen.add(f.check_id)
            text = sec["bullets"].get(f.check_id)
            if text:
                _add_bullet(doc, "", text, branding)
            else:
                _add_bullet(doc, f.title, f.description, branding)
        _add_paragraph(doc, sec["closer"], color=branding.color_body)

    return next_num


def _risk_classification(doc, branding, host: str, grade: str) -> None:
    _add_section_heading(doc, "3. Kockázati besorolás", branding)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    level = phrases.grade_to_risk_level(grade)
    levels = [
        ("ALACSONY", "16a34a", level == "low"),
        ("KÖZEPES", "ca8a04", level == "medium"),
        ("MAGAS", "B00020", level == "high"),
    ]
    for i, (label, color, active) in enumerate(levels):
        prefix = "  ●  " if active else "  ○  "
        _add_run(p, prefix, color=color, size=12)
        _add_run(p, label, bold=active, color=color if active else branding.color_meta, size=11)
        if i < len(levels) - 1:
            _add_run(p, "       ", color=branding.color_meta)

    _add_paragraph(doc, phrases.RISK_CLASSIFICATION[level].format(host=host), color=branding.color_body)


def _offers(doc, branding, cfg: Config, *, price_offer1: str, duration_offer1: str,
            price_offer2: str, duration_offer2: str) -> None:
    _add_section_heading(doc, "4. Javasolt megoldások", branding)

    # Offer 1
    o1 = phrases.OFFER_1
    p = doc.add_paragraph()
    _add_run(p, o1["title"], bold=True, color=branding.color_primary, size=12)
    sub = doc.add_paragraph()
    r = sub.add_run(o1["subtitle"])
    r.italic = True
    r.font.color.rgb = _rgb(branding.color_meta)
    r.font.size = Pt(10)

    _add_paragraph(doc, o1["intro"], color=branding.color_body)
    for b in o1["bullets"]:
        _add_bullet(doc, "", b, branding)
    _add_paragraph(doc, f"Eredmény: {o1['result']}", color=branding.color_body, bold=False)

    meta = doc.add_paragraph()
    _add_run(meta, "Időtartam: ", bold=True, color=branding.color_body)
    _add_run(meta, duration_offer1, color=branding.color_body)
    if price_offer1:
        meta2 = doc.add_paragraph()
        _add_run(meta2, "Irányár: ", bold=True, color=branding.color_body)
        _add_run(meta2, price_offer1, color=branding.color_critical)

    _add_paragraph(doc, o1["footer"], color=branding.color_meta)

    # Offer 2
    o2 = phrases.OFFER_2
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    _add_run(p, o2["title"], bold=True, color=branding.color_primary, size=12)
    sub = doc.add_paragraph()
    r = sub.add_run(o2["subtitle"])
    r.italic = True
    r.font.color.rgb = _rgb(branding.color_meta)
    r.font.size = Pt(10)
    for b in o2["bullets"]:
        _add_bullet(doc, "", b, branding)
    _add_paragraph(doc, f"Miért merül fel ez opcióként: {o2['why']}", color=branding.color_body)
    meta = doc.add_paragraph()
    _add_run(meta, "Időtartam: ", bold=True, color=branding.color_body)
    _add_run(meta, duration_offer2, color=branding.color_body)
    meta2 = doc.add_paragraph()
    _add_run(meta2, "Irányár: ", bold=True, color=branding.color_body)
    _add_run(meta2, price_offer2, color=branding.color_critical)


def _next_step(doc, branding, host: str) -> None:
    _add_section_heading(doc, "5. Javaslat és következő lépés", branding)
    text = phrases.NEXT_STEP_TEMPLATE.format(host=host)
    for para in text.split("\n\n"):
        _add_paragraph(doc, para, color=branding.color_body)


def _footer_block(doc, branding) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, f"{branding.email}    ·    {branding.company_name}    ·    {branding.website}",
             color=branding.color_meta, size=8)


def build(
    *,
    target: str,
    results: list[CheckerResult],
    output_path: str,
    cfg: Config,
    industry: str | None = None,
    compliance: str | None = None,
    price_offer1: str | None = None,
    duration_offer1: str | None = None,
    price_offer2: str | None = None,
    duration_offer2: str | None = None,
    scan_source: str | None = None,
) -> str:
    branding = cfg.branding
    defaults = cfg.defaults

    industry = industry or defaults.industry
    compliance = compliance or defaults.compliance
    price_offer1 = price_offer1 if price_offer1 is not None else defaults.price_offer1
    duration_offer1 = duration_offer1 or defaults.duration_offer1
    price_offer2 = price_offer2 if price_offer2 is not None else defaults.price_offer2
    duration_offer2 = duration_offer2 or defaults.duration_offer2
    scan_source = scan_source or defaults.scan_source

    parsed = urlparse(target if "://" in target else f"https://{target}")
    host = parsed.hostname or target
    target_url = f"{parsed.scheme}://{parsed.netloc}/" if parsed.scheme else target

    score, grade = compute_score(results)
    info = _classify_findings(results)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = _rgb(branding.color_body)

    _add_title_block(doc, branding)
    now = datetime.now()
    months_hu = ["január", "február", "március", "április", "május", "június",
                 "július", "augusztus", "szeptember", "október", "november", "december"]
    scan_date = f"{now.year}. {months_hu[now.month - 1]} {now.day}."
    _add_info_table(
        doc, branding,
        target_url=target_url,
        scan_date=scan_date,
        grade=grade,
        scan_source=scan_source,
    )

    _executive_summary(doc, branding, host=host, grade=grade, info=info,
                       industry=industry, compliance=compliance)

    _add_section_heading(doc, "2. Technikai megállapítások", branding)
    _section_2_headers(doc, branding, info)
    _section_2_adaptive(doc, branding, info, results)

    _risk_classification(doc, branding, host, grade)
    _offers(doc, branding, cfg,
            price_offer1=price_offer1, duration_offer1=duration_offer1,
            price_offer2=price_offer2, duration_offer2=duration_offer2)
    _next_step(doc, branding, host)

    _footer_block(doc, branding)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
